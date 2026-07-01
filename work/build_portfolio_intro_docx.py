from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\ProductSenseAI")
MD_PATH = ROOT / "docs" / "ProductSenseAI_Portfolio_Introduction.md"
DOCX_PATH = ROOT / "docs" / "ProductSenseAI_Portfolio_Introduction.docx"

BODY_FONT = "Calibri"
EAST_ASIA_FONT = "Microsoft YaHei"
HEADING_BLUE = RGBColor(46, 116, 181)
HEADING_DARK = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 90, 90)
TABLE_FILL = "F2F4F7"


def apply_run_font(run, size=None, bold=None, color=None):
    run.font.name = BODY_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = color


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_specs = [
        ("Heading 1", 16, HEADING_BLUE, 16, 8),
        ("Heading 2", 13, HEADING_BLUE, 12, 6),
        ("Heading 3", 12, HEADING_DARK, 8, 4),
    ]
    for name, size, color, before, after in heading_specs:
        style = styles[name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text.strip())
    apply_run_font(run, size=10, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def clean_inline(text):
    return text.replace("`", "").strip()


def add_paragraph_with_bold(doc, text, style=None):
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.line_spacing = 1.10
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            apply_run_font(run, bold=True)
        else:
            run = paragraph.add_run(part)
            apply_run_font(run)
    return paragraph


def add_markdown_table(doc, table_lines):
    rows = []
    for line in table_lines:
        if re.match(r"^\s*\|?\s*:?-{3,}:?\s*\|", line):
            continue
        cells = [clean_inline(cell) for cell in line.strip().strip("|").split("|")]
        if len(cells) > 1:
            rows.append(cells)

    if not rows:
        return

    column_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=column_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False

    for row_index, row in enumerate(rows):
        for col_index in range(column_count):
            cell = table.cell(row_index, col_index)
            value = row[col_index] if col_index < len(row) else ""
            is_header = row_index == 0 or col_index == 0
            set_cell_text(cell, value, bold=is_header)
            if is_header:
                set_cell_shading(cell, TABLE_FILL)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def add_title(doc, title):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(24)
    paragraph.paragraph_format.space_after = Pt(10)
    run = paragraph.add_run(title)
    apply_run_font(run, size=22, bold=True, color=HEADING_DARK)


def add_footer(doc):
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("ProductSense AI | AI Product Manager Portfolio MVP")
        apply_run_font(run, size=8.5, color=MUTED)


def build():
    markdown = MD_PATH.read_text(encoding="utf-8")
    lines = markdown.splitlines()

    doc = Document()
    configure_document(doc)

    index = 0
    if lines and lines[0].startswith("# "):
        add_title(doc, lines[0][2:].strip())
        index = 1

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if not stripped:
            index += 1
            continue

        if stripped.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index])
                index += 1
            add_markdown_table(doc, table_lines)
            continue

        if stripped.startswith("### "):
            doc.add_heading(clean_inline(stripped[4:]), level=3)
        elif stripped.startswith("## "):
            doc.add_heading(clean_inline(stripped[3:]), level=1)
        elif stripped.startswith(">"):
            paragraph = add_paragraph_with_bold(doc, clean_inline(stripped.lstrip(">").strip()))
            paragraph.paragraph_format.left_indent = Inches(0.25)
            for run in paragraph.runs:
                run.font.color.rgb = MUTED
        elif stripped.startswith("- "):
            add_paragraph_with_bold(doc, clean_inline(stripped[2:]), style="List Bullet")
        elif re.match(r"^\d+\.\s+", stripped):
            add_paragraph_with_bold(doc, clean_inline(re.sub(r"^\d+\.\s+", "", stripped)), style="List Number")
        else:
            add_paragraph_with_bold(doc, stripped)

        index += 1

    add_footer(doc)
    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build()

